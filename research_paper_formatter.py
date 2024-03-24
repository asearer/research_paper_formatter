import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document

def create_blank_template(output_file):
    document = Document()

    # Add title
    title = document.add_heading(level=1)
    title_run = title.add_run("Title")
    title_run.bold = True

    # Add author
    author = document.add_paragraph("Author Name")
    author.alignment = 1  # Center alignment

    # Add affiliation
    affiliation = document.add_paragraph("Affiliation")
    affiliation.alignment = 1  # Center alignment

    # Add abstract
    abstract_heading = document.add_heading(level=2)
    abstract_heading_run = abstract_heading.add_run("Abstract")
    abstract_heading_run.bold = True

    abstract = document.add_paragraph("Abstract content goes here.")

    # Add main content section
    document.add_page_break()
    main_heading = document.add_heading(level=1)
    main_heading_run = main_heading.add_run("Main Content")
    main_heading_run.bold = True

    # Add introduction
    document.add_heading("Introduction", level=2)

    # Add body paragraphs
    document.add_paragraph("Body paragraph 1")
    document.add_paragraph("Body paragraph 2")
    document.add_paragraph("Body paragraph 3")

    # Add conclusion
    document.add_heading("Conclusion", level=2)

    # Save the document
    document.save(output_file)
    messagebox.showinfo("Success", "Blank template created successfully.")

def browse_output_file_template():
    output_file = filedialog.asksaveasfilename(defaultextension=".docx")
    if output_file:
        output_entry_template.delete(0, tk.END)
        output_entry_template.insert(0, output_file)

def create_template():
    output_file = output_entry_template.get()
    if not output_file:
        messagebox.showerror("Error", "Please select an output file.")
        return
    create_blank_template(output_file)

# Create the main window
root = tk.Tk()
root.title("Research Paper Formatter")

# Create output file selection widgets for template
output_label_template = tk.Label(root, text="Output File:")
output_label_template.grid(row=0, column=0, padx=5, pady=5)
output_entry_template = tk.Entry(root, width=50)
output_entry_template.grid(row=0, column=1, padx=5, pady=5)
output_button_template = tk.Button(root, text="Browse", command=browse_output_file_template)
output_button_template.grid(row=0, column=2, padx=5, pady=5)

# Create create template button
create_template_button = tk.Button(root, text="Create Template", command=create_template)
create_template_button.grid(row=1, column=1, padx=5, pady=5)

# Run the main event loop
root.mainloop()
